# This allows you to find all the chunks by a specific meeting date. 

import os
import weaviate
from dotenv import load_dotenv
from docx import Document

# Load environment variables from .env
load_dotenv()

# Initialize Weaviate client
WEAVIATE_URL = os.getenv("WEAVIATE_URL")
WEAVIATE_API_KEY = os.getenv("WEAVIATE_API_KEY")
client = weaviate.Client(
    url=WEAVIATE_URL,
    auth_client_secret=weaviate.AuthApiKey(api_key=WEAVIATE_API_KEY)
)

def fetch_documents_by_date_and_export_to_word(date):
    """
    Fetch documents from Weaviate filtered by a specific date and export metadata, including source_document, to a Word document.

    Args:
        date (str): The date to filter by (YYYY-MM-DD format).
    """
    query = f"""
    {{
      Get {{
        MeetingDocument(where: {{
          path: ["meeting_date"],
          operator: Equal,
          valueString: "{date}"
        }}) {{
          content
          meeting_date
          meeting_type
          file_type
          chunk_index
          source_document
        }}
      }}
    }}
    """
    try:
        print(f"Querying Weaviate for documents on {date}...")
        response = client.query.raw(query)
        documents = response.get("data", {}).get("Get", {}).get("MeetingDocument", [])
        
        if not documents:
            print(f"No documents found for the date: {date}.")
            return

        print(f"\nRetrieved Documents for {date}:")
        for doc in documents:
            print(f"- Chunk Index: {doc.get('chunk_index', 'N/A')}")
            print(f"  Meeting Date: {doc.get('meeting_date', 'N/A')}")
            print(f"  Meeting Type: {doc.get('meeting_type', 'N/A')}")
            print(f"  File Type: {doc.get('file_type', 'N/A')}")
            print(f"  Source Document: {doc.get('source_document', 'N/A')}")
            print(f"  Content Preview: {doc.get('content', 'N/A')[:100]}...")
            print()

        # Export metadata to Word
        print(f"Exporting metadata for {date} to Word document...")
        doc = Document()
        doc.add_heading(f'Document Metadata for {date}', level=1)

        for doc_data in documents:
            doc.add_heading(f"Chunk Index: {doc_data.get('chunk_index', 'N/A')}", level=2)
            doc.add_paragraph(f"Meeting Date: {doc_data.get('meeting_date', 'N/A')}")
            doc.add_paragraph(f"Meeting Type: {doc_data.get('meeting_type', 'N/A')}")
            doc.add_paragraph(f"File Type: {doc_data.get('file_type', 'N/A')}")
            doc.add_paragraph(f"Source Document: {doc_data.get('source_document', 'N/A')}")
            doc.add_paragraph(f"Content Preview: {doc_data.get('content', 'N/A')}")
            doc.add_paragraph("\n")

        word_file_path = f"Weaviate_Metadata_List_{date}.docx"
        doc.save(word_file_path)
        print(f"Metadata exported to {word_file_path} successfully.")

    except Exception as e:
        print(f"Error querying Weaviate: {e}")

if __name__ == "__main__":
    # Filter by specific date (YYYY-MM-DD format)
    specific_date = "2000-10-27"
    fetch_documents_by_date_and_export_to_word(specific_date)